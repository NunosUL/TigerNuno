"""
generate_setup_guide.py
-----------------------
Single script to create or regenerate TigerChat_Setup_Guide.docx.

Run:
    python generate_setup_guide.py

The document is always rebuilt from scratch from the content in this file —
this script is the single source of truth for the published guide.

Images are sourced from setup_guide_assets/ (image1.png, image2.png, image3.png).
Run this script once after cloning; the assets folder is committed to git.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH  = os.path.join(os.path.dirname(__file__), "TigerChat_Setup_Guide.docx")
ASSETS_DIR   = os.path.join(os.path.dirname(__file__), "setup_guide_assets")


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Apply a solid background fill to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    # Remove existing shd if present
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_heading1_color(para, hex_color="2F81F7"):
    """Apply color to all runs in a Heading 1 paragraph."""
    for run in para.runs:
        run.font.color.rgb = RGBColor.from_string(hex_color)


# ── Document-level helpers ────────────────────────────────────────────────────

def _heading1(doc, text: str):
    p = doc.add_heading(text, level=1)
    _set_heading1_color(p)
    return p


def _heading2(doc, text: str):
    return doc.add_heading(text, level=2)


def _para(doc, text: str = ""):
    p = doc.add_paragraph(text, style="Normal")
    return p


def _bold_para(doc, label: str, rest: str = ""):
    p = doc.add_paragraph(style="Normal")
    r = p.add_run(label)
    r.bold = True
    if rest:
        p.add_run(rest)
    return p


def _bullet(doc, text: str):
    doc.add_paragraph(text, style="List Bullet")


def _numbered(doc, label: str, rest: str = ""):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(label)
    r.bold = True
    if rest:
        p.add_run(" — " + rest)


def _code_block(doc, lines: list[str]):
    """Render a code block as a single-cell table with light gray background."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_bg(cell, "F6F8FA")
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x17, 0x5B, 0xC2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)


def _callout(doc, text: str, bg_hex: str, text_hex: str = "000000"):
    """Single-cell colored callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_bg(cell, bg_hex)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(text_hex)
    run.font.size = Pt(10)


def _spacer(doc):
    p = doc.add_paragraph("")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


def _image(doc, filename: str, width_inches: float = 5.52):
    path = os.path.join(ASSETS_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
    else:
        _para(doc, f"[Image: {filename} — run generate_setup_guide.py to embed]")


def _styled_table(doc, headers: list[str], rows: list[list[str]],
                  header_hex: str = "2F81F7",
                  font_size: float = 9.5):
    """
    Add a formatted table.
    - Header row: header_hex background, white bold text.
    - Data rows: alternating #F6F8FA / #FFFFFF.
    """
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"

    # Header
    hdr_cells = table.rows[0].cells
    for ci, h in enumerate(headers):
        _set_cell_bg(hdr_cells[ci], header_hex)
        hdr_cells[ci].paragraphs[0].clear()
        p = hdr_cells[ci].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(font_size)

    # Data rows
    alt = ["F6F8FA", "FFFFFF"]
    for ri, row_data in enumerate(rows):
        bg = alt[ri % 2]
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            _set_cell_bg(cells[ci], bg)
            cells[ci].paragraphs[0].clear()
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(font_size)

    return table


# ── Document content ──────────────────────────────────────────────────────────

def build_document(doc):

    # ── Cover page ────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TigerChat")
    r.bold      = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x2F, 0x81, 0xF7)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Setup & Architecture Guide")
    rs.font.size = Pt(18)
    rs.font.color.rgb = RGBColor(0x24, 0x29, 0x2F)

    _spacer(doc)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = tagline.add_run(
        "Wiki · Code · Test Cases · Work Items · Commit Diffs"
    )
    rt.font.size = Pt(14)
    rt.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    tech = doc.add_paragraph()
    tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rtech = tech.add_run("Azure AI Search · Azure OpenAI GPT-4o · FastAPI")
    rtech.font.size = Pt(12)
    rtech.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()

    # ── 1. Overview ───────────────────────────────────────────────────────────
    _heading1(doc, "1. Overview")
    _para(doc,
        "TigerChat is an internal RAG-powered chat application that lets teams ask "
        "natural-language questions about their Azure DevOps content. Answers are "
        "grounded in five source types, retrieved via hybrid vector + BM25 search, "
        "and answered by GPT-4o with a live confidence signal.")

    _spacer(doc)
    _image(doc, "image1.png", width_inches=5.52)   # Ingestion pipeline diagram
    _spacer(doc)
    _image(doc, "image2.png", width_inches=5.52)   # Query pipeline diagram
    _spacer(doc)

    _styled_table(doc,
        ["Source", "What gets indexed"],
        [
            ["📄 Wiki pages",
             "All pages across all wikis, recursively, as Markdown with discussion comments"],
            ["💻 Source code",
             ".cs .csproj .razor .ts .js .json .xml .config and more (up to 100 KB per file)"],
            ["🧪 Test cases",
             "All test cases with steps, expected results, shared steps resolved, "
             "and discussion threads"],
            ["📋 Work Items",
             "Epics, Features, User Stories, Tasks, and Bugs — descriptions, "
             "acceptance criteria, linked items, comments, and code/config attachments. "
             "Bugs/Tasks also resolve linked Git commits for diff coverage."],
            ["🔀 Commit Diffs",
             "Unified before/after file diffs for every commit linked to a Bug or Task, "
             "including NuGet 'Package | Before | After' change tables"],
        ],
        header_hex="2F81F7",
    )

    # ── 2. Prerequisites ──────────────────────────────────────────────────────
    _heading1(doc, "2. Prerequisites")

    _heading2(doc, "Azure Resources")
    for item in [
        "Azure AI Search — Standard S1 or higher (required for Semantic Reranker)",
        "Azure OpenAI — with deployments: text-embedding-3-large, gpt-4o",
        "Azure Blob Storage — container for ingestion snapshots and hash manifest",
        "Azure DevOps — project with Wiki, Git repos, Test Plans, and Work Items enabled",
    ]:
        _bullet(doc, item)
    _spacer(doc)

    _heading2(doc, "PAT Token Scopes")
    _para(doc, "Your Azure DevOps Personal Access Token must have ALL four scopes enabled:")
    _spacer(doc)
    _callout(doc,
        "⚠️  Important — Your PAT must have ALL four scopes enabled: "
        "Wiki (Read), Code (Read), Test Management (Read), Work Items (Read). "
        "Missing any scope will cause silent failures for that source type during ingestion.",
        bg_hex="FFB066",
    )
    _spacer(doc)
    _styled_table(doc,
        ["Scope", "Level", "Required for"],
        [
            ["Wiki",            "Read", "Crawling wiki pages and discussion comments"],
            ["Code",            "Read", "Enumerating repos and downloading source files"],
            ["Test Management", "Read", "Fetching test plans, suites, test cases, and steps"],
            ["Work Items",      "Read", "WI details, comments, attachments, and commit link resolution"],
        ],
        header_hex="2F81F7",
    )
    _spacer(doc)

    _heading2(doc, "Software")
    _bullet(doc, "Python 3.11+")
    _bullet(doc, "pip install -r requirements.txt")

    # ── 3. Environment Variables ───────────────────────────────────────────────
    _heading1(doc, "3. Environment Variables")
    _para(doc, "Copy .env.example to .env and fill in all required values.")
    _spacer(doc)
    _styled_table(doc,
        ["Variable", "Description", "Default"],
        [
            ["AZURE_DEVOPS_PAT",                   "Personal Access Token (all four scopes above)",         "(required)"],
            ["DEVOPS_ORG",                         "Azure DevOps organisation name",                       "(required)"],
            ["DEVOPS_PROJECT",                     "Project name",                                         "(required)"],
            ["AZURE_STORAGE_CONNECTION_STRING",    "Blob Storage connection string",                       "(required)"],
            ["AZURE_STORAGE_CONTAINER",            "Blob container name for snapshots and manifest",       "wiki-crawl"],
            ["AZURE_OPENAI_ENDPOINT",              "Azure OpenAI resource URL",                            "(required)"],
            ["AZURE_OPENAI_API_KEY",               "Azure OpenAI API key",                                 "(required)"],
            ["AZURE_OPENAI_EMBEDDING_DEPLOYMENT",  "Embedding model deployment name",                      "text-embedding-3-large"],
            ["AZURE_OPENAI_API_VERSION",           "API version",                                          "2024-02-01"],
            ["AZURE_OPENAI_CHAT_DEPLOYMENT",       "Chat model deployment name",                           "gpt-4o"],
            ["AZURE_SEARCH_ENDPOINT",              "Azure AI Search endpoint URL",                         "(required)"],
            ["AZURE_SEARCH_API_KEY",               "Azure AI Search API key",                              "(required)"],
            ["AZURE_SEARCH_INDEX_NAME",            "Search index name",                                    "wiki-index"],
            ["CRAWL_WIKI",                         "Include Wiki pages in ingestion",                      "true"],
            ["CRAWL_CODE",                         "Include source code in ingestion",                     "true"],
            ["CRAWL_TESTS",                        "Include test cases in ingestion",                      "true"],
            ["CRAWL_WORK_ITEMS",                   "Include Work Items (Epics, Features, Stories, Tasks, Bugs)", "true"],
            ["CRAWL_COMMIT_DIFFS",                 "Include commit diffs linked to Bugs/Tasks",             "true"],
            ["CRAWL_WI_ATTACHMENTS",               "Download and embed code/config attachments from Work Items", "true"],
            ["AZURE_SEARCH_SEMANTIC_ENABLED",      "Enable Semantic Reranker (S1+ tier required)",         "true"],
            ["BOOST_SOURCE_TYPE",                  "Boost source: wiki / code / test / workitem / empty",  ""],
            ["EMBED_INTER_BATCH_DELAY_S",          "Pause between embedding batches (rate-limit tuning)",  "0.5"],
        ],
        header_hex="2F81F7",
    )

    # ── 4. Installation ───────────────────────────────────────────────────────
    _heading1(doc, "4. Installation")
    _code_block(doc, [
        "git clone <repository-url>",
        "cd TigerNuno",
        "pip install -r requirements.txt",
        "cp .env.example .env",
        "# Fill in all required values in .env",
        "uvicorn app:app --reload",
    ])
    _spacer(doc)
    _callout(doc,
        "✅  Ready — Open http://localhost:8000 to access TigerChat. "
        "The application serves the landing page at /, the Ingestion UI at /ingest, "
        "and the Chat interface at /chat.",
        bg_hex="67E178",
        text_hex="1A5C2A",
    )

    # ── 5. Ingestion Pipeline ─────────────────────────────────────────────────
    _heading1(doc, "5. Ingestion Pipeline")

    _heading2(doc, "Overview")
    _para(doc,
        "Transforms raw Azure DevOps content into searchable vector chunks in five steps. "
        "Controlled from the Ingestion UI at /ingest with per-source toggles, repo picker, "
        "test plan/suite/test-case picker, and area path picker for Work Items. "
        "Live SSE progress streaming shows per-step counts in real time.")

    _heading2(doc, "Step 1 — Crawl")
    _para(doc, "Each source type is crawled via the Azure DevOps REST API:")
    _bullet(doc,
        "Wiki: Lists all pages recursively (recursionLevel=full), fetches Markdown "
        "content, crawls discussion comments per page.")
    _bullet(doc,
        "Source Code: Enumerates all Git repos (or selected repos via repo picker), "
        "walks file tree (paginated via x-ms-continuationtoken), downloads text files "
        "up to 100 KB. Extensions: .cs .csproj .sln .config .resx .razor .html .js "
        ".ts .json .xml .xslt. Skips disabled repos and repos with no default branch.")
    _bullet(doc,
        "Test Cases: Lists test plans → suites → work item IDs → fetches full work "
        "items with steps, expected results, shared steps resolved, and discussion "
        "threads. Requires Work Items (Read) PAT scope.")
    _bullet(doc,
        "Work Items: Queries each type (Epic, Feature, User Story, Task, Bug) "
        "separately via WIQL to stay under the 20,000 item limit. Fetches full "
        "details: description, acceptance criteria, repro steps, linked items, "
        "discussion comments, and code/config attachments. Bugs and Tasks have "
        "ArtifactLink relations resolved to real Git commits.")
    _bullet(doc,
        "Commit Diffs: For each unique commit linked to a Bug or Task, fetches the "
        "unified before/after diff for every text file changed (up to 20 files, "
        "30 KB per file). .csproj files produce a Package | Before | After diff "
        "table. Stored at path /commit-diffs/{sha}.")

    _heading2(doc, "Step 2 — Chunk")
    _para(doc,
        "Uses LangChain RecursiveCharacterTextSplitter (chunk_size=1000, overlap=150). "
        "Each chunk stores: id, text, path, url, crawled_at, "
        "source_type (wiki / code / test / workitem / commit-diff), "
        "source_tags ([\"wiki\"] / [\"code\"] / [\"test\"] / [\"workitem\"] / [\"commit-diff\"]).")

    _heading2(doc, "Step 3 — Embed")
    _para(doc,
        "Azure OpenAI text-embedding-3-large, 3072 dimensions, batched in groups of 200 "
        "with 3 concurrent workers. EMBED_INTER_BATCH_DELAY_S prevents rate limit errors.")

    _heading2(doc, "Step 4 — Index Schema")
    _styled_table(doc,
        ["Field", "Type", "Purpose"],
        [
            ["id",          "String (key)",                    "Unique chunk identifier"],
            ["text",        "SearchableString",                "Chunk content for BM25 and Semantic Reranker"],
            ["path",        "String (filterable)",             "Source path — /repos/…, /work-items/301615, /commit-diffs/{sha}"],
            ["url",         "String",                          "Full Azure DevOps URL for clickable citations"],
            ["crawled_at",  "String (filterable)",             "Crawl timestamp for change detection"],
            ["source_type", "String (filterable, facetable)",  "wiki / code / test / workitem / commit-diff"],
            ["source_tags", "Collection(String) (filterable)", "[\"wiki\"] / [\"code\"] / [\"test\"] / [\"workitem\"] / [\"commit-diff\"]"],
            ["embedding",   "Collection(Single) — HNSW",       "3072-dim vector, cosine similarity"],
        ],
        header_hex="3FB950",
    )
    _spacer(doc)
    _para(doc,
        "The index defines a SemanticConfiguration (\"tigerchat-semantic\") and a "
        "ScoringProfile (\"source-boost\") with a TagScoringFunction applying a 3× "
        "boost to matching source_tags.")

    _heading2(doc, "Step 5 — Upload")
    _para(doc,
        "Chunks are merged/upserted into Azure AI Search in batches. Change detection "
        "uses an MD5 hash manifest stored in Blob Storage (rag_hashes.json) — only "
        "content whose hash changed is re-processed. The raw crawl snapshot is stored "
        "as rag_{ORG}_{PROJECT}.jsonl.")
    _spacer(doc)
    _styled_table(doc,
        ["Blob", "Purpose"],
        [
            ["rag_{ORG}_{PROJECT}.jsonl", "Raw crawl snapshot — used by crawl=OFF re-process mode"],
            ["rag_hashes.json",           "MD5 hash manifest — drives incremental change detection"],
        ],
        header_hex="3FB950",
    )
    _spacer(doc)
    _callout(doc,
        "ℹ️  First run: both blobs are absent. The full pipeline runs and creates them. "
        "Subsequent runs skip unchanged content. crawl=OFF mode requires the snapshot "
        "blob to already exist (run crawl=ON at least once first).",
        bg_hex="D0E8FF",
        text_hex="0A3D62",
    )

    # ── 6. Ingestion UI Settings ───────────────────────────────────────────────
    _heading1(doc, "6. Ingestion UI Settings")
    _bullet(doc,
        "Enable Crawl toggle: when off, skips crawling and re-processes the last "
        "snapshot (useful for re-embedding without a fresh crawl).")
    _bullet(doc,
        "Source checkboxes: Wiki / Source Code / Test Cases / Work Items / Commit "
        "Diffs — select which sources to include in this run.")
    _bullet(doc, "Repo picker: select specific Git repositories to include (or all).")
    _bullet(doc, "Test plan / suite / test-case picker: narrow test case ingestion.")
    _bullet(doc, "Area path picker: limit Work Items ingestion to specific area paths.")

    # ── 7. Chat RAG Pipeline ──────────────────────────────────────────────────
    _heading1(doc, "7. Chat RAG Pipeline")

    _heading2(doc, "Overview")
    _para(doc,
        "A five-step pipeline turns a natural-language question into a grounded "
        "answer with clickable citations and a confidence signal.")

    _heading2(doc, "Step 1 — Embed Query")
    _para(doc, "Question embedded using text-embedding-3-large → 3072-dimensional vector.")

    _heading2(doc, "Step 2 — Hybrid Search + Semantic Reranker")
    _para(doc,
        "Vector similarity (cosine, HNSW) + BM25 run in parallel. Results merged via "
        "RRF: score = 1/(rank_vector+60) + 1/(rank_bm25+60). Top 15 re-scored by "
        "Semantic Reranker (0–4). Chunks with reranker_score < 1.0 are dropped.")

    _heading2(doc, "Step 3 — Commit Diff Injection")
    _para(doc,
        "Retrieved Work Item chunks are scanned for *sha:40-char-hex* markers. "
        "For each unique SHA, the corresponding /commit-diffs/{sha} document is "
        "fetched and injected into context — enabling PFQ-style answers about "
        "code changes, NuGet upgrades, and file modifications.")

    _heading2(doc, "Step 4 — Context Builder")
    _para(doc,
        "Chunks ranked by reranker_score, deduplicated by ID, soft-capped at "
        "4,000 chars per source path. Each chunk prefixed with its source type label. "
        "Concatenated up to ~16k chars context limit.")

    _heading2(doc, "Step 5 — LLM Generation")
    _para(doc,
        "GPT-4o (temperature=0.2, max=1024 tokens). System prompt instructs the model "
        "to answer using only provided context, list test cases with IDs and steps, "
        "reference file paths for code, and not to append a Sources list.")

    _heading2(doc, "Step 6 — Format Response")
    _para(doc,
        "Answer returned with clickable citations linking to Azure DevOps, "
        "a confidence signal (High / Medium / Low), and markdown tables rendered as "
        "styled HTML tables with a copy-to-clipboard button.")

    # ── 8. Scoring & Confidence ────────────────────────────────────────────────
    _heading1(doc, "8. Scoring & Confidence")

    _image(doc, "image3.png", width_inches=4.24)   # Scoring visualization
    _spacer(doc)

    _heading2(doc, "Score Metrics")
    _styled_table(doc,
        ["Metric", "Scale", "Source"],
        [
            ["@search.score",          "0.01 – 0.06", "RRF fusion score from hybrid search"],
            ["@search.reranker_score", "0 – 4",       "Neural semantic reranker (S1+ tier only)"],
        ],
        header_hex="A371F7",
    )
    _spacer(doc)

    _heading2(doc, "Confidence Gates")
    _styled_table(doc,
        ["Gate", "Condition", "Signal"],
        [
            ["Gate 1", "No chunks retrieved",                       "🔴 LOW"],
            ["Gate 2", "LLM answer contains uncertainty phrase",     "🔴 LOW"],
            ["Gate 3 (semantic on)", "reranker_score ≥ 2.5 AND chunks ≥ 3", "🟢 HIGH"],
            ["Gate 3 (semantic on)", "reranker_score ≥ 1.5 AND chunks ≥ 2", "🟡 MEDIUM"],
            ["Gate 3 (semantic on)", "Anything else",                "🔴 LOW"],
            ["Gate 3 (fallback)",    "@search.score > 0.03 AND chunks ≥ 4", "🟢 HIGH"],
            ["Gate 3 (fallback)",    "@search.score > 0.01 AND chunks ≥ 2", "🟡 MEDIUM"],
        ],
        header_hex="F0883E",
    )
    _spacer(doc)

    _heading2(doc, "Reranker Score Zones")
    _styled_table(doc,
        ["Score Range", "Zone", "Confidence"],
        [
            ["< 1.0",    "Below min-score floor — filtered out before context", "—"],
            ["1.0 – 1.5", "Low relevance",      "🔴 LOW"],
            ["1.5 – 2.5", "Moderate relevance", "🟡 MEDIUM (≥ 2 chunks)"],
            ["2.5 – 4.0", "High relevance",      "🟢 HIGH (≥ 3 chunks)"],
        ],
        header_hex="24292F",
    )

    # ── 9. Application Routes ──────────────────────────────────────────────────
    _heading1(doc, "9. Application Routes")
    _styled_table(doc,
        ["Route", "Method", "Description"],
        [
            ["/",                          "GET",      "Landing page"],
            ["/ingest",                    "GET",      "Ingestion UI with live SSE progress and pickers"],
            ["/chat",                      "GET",      "Chat interface with 5-source accordion example questions"],
            ["/about",                     "GET",      "How Ingestion Works"],
            ["/about/diffs",               "GET",      "How Commit Diff ingestion works"],
            ["/chat/about",                "GET",      "How Chat Works (RAG pipeline)"],
            ["/chat/about/scoring",        "GET",      "Scoring & Confidence detail page"],
            ["/synergies",                 "GET",      "Pipeline relationship diagram"],
            ["/api/chat",                  "POST",     "Accepts {question, source_types?}, returns {answer, sources, confidence}"],
            ["/api/chat/stream",           "POST SSE", "Streams per-step progress events then the final answer"],
            ["/ingest/stream",             "GET SSE",  "Streams real-time ingestion pipeline progress"],
            ["/api/repos",                 "GET",      "Lists available Git repositories for the repo picker"],
            ["/api/test-plans",            "GET",      "Lists all test plans"],
            ["/api/test-suites",           "GET",      "Lists all suites across all plans"],
            ["/api/areas",                 "GET",      "Full area path tree with depth info"],
            ["/api/areas/counts",          "GET",      "Work item counts per type for selected area paths"],
            ["/api/index/stats",           "GET",      "Document counts per source type in the search index"],
            ["/api/debug/tc/{id}",         "GET",      "Indexed chunks and raw DevOps fields for a test case"],
            ["/api/debug/wi/{id}",         "GET",      "Indexed chunks and ArtifactLink resolution for a Work Item"],
            ["/api/debug/wi/{id}/dev-info","GET",      "Step-by-step commit resolution trace for a Work Item"],
        ],
        header_hex="1AAC8F",
    )

    # ── 10. Key Design Decisions ───────────────────────────────────────────────
    _heading1(doc, "10. Key Design Decisions")
    decisions = [
        ("Hybrid search (vector + BM25)",
         "Handles both semantic paraphrasing and exact term queries, outperforming pure vector search."),
        ("Semantic Reranker (0–4 score)",
         "Neural cross-encoder re-scores RRF results, replacing brittle RRF score thresholds."),
        ("Source type labels in context",
         "[Wiki/Source Code/Test Case/Work Item/Commit Diff] so GPT-4o applies different answer styles per content type."),
        ("Commit Diff injection via sha: markers",
         "Work Item chunks embed *sha:* markers at ingest; at query time the matching diff chunks are injected automatically for PFQ-quality change reports."),
        ("Source type filter (source_types parameter)",
         "The API accepts an optional source_types list to restrict the Azure AI Search OData filter and pinning logic, keeping test-focused questions free of unrelated work items and source code."),
        ("Tag-based scoring profile",
         "Enables boosting a preferred source type at runtime via BOOST_SOURCE_TYPE without re-indexing."),
        ("Min-score filter (reranker_score < 1.0)",
         "Prevents weakly-matched chunks from degrading answer quality before they reach GPT-4o."),
        ("create_or_update_index",
         "Schema changes can be applied to existing indexes without data loss."),
        ("MD5 hash manifest (rag_hashes.json)",
         "Change detection prevents re-processing unchanged content across ingestion runs."),
        ("WIQL per-type queries",
         "Querying each work item type separatley stays under the 20,000-item WIQL result limit."),
    ]
    for label, desc in decisions:
        _numbered(doc, label, desc)


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    doc = Document()

    # Margins matching the Copy.docx (~1 inch on all sides)
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    build_document(doc)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}  ({os.path.getsize(OUTPUT_PATH) // 1024} KB)")
