"""
generate_pfq_doc.py
-------------------
Generates a professional IT audit (PFQ) Word document for code-change analysis.

Layout per change entry
───────────────────────
┌─────────────────────────────────────────────────────────────────────────┐
│  Consolas · grey bg  │  /Path/To/File.cs                                │
├────────────┬────────────────────────────────────────────────────────────┤
│  BEFORE    │  Plain-text description (optional inline code block)        │
│  (bold red)│                                                             │
├────────────┼────────────────────────────────────────────────────────────┤
│  AFTER     │  Plain-text description                                     │
│  (bold grn)│  ╔══════════════════════════════════════════════════════╗  │
│            │  ║ Consolas 9pt · grey bg · box border                 ║  │
│            │  ║ if (x == Guid.Empty) { ... }                        ║  │
│            │  ╚══════════════════════════════════════════════════════╝  │
└────────────┴────────────────────────────────────────────────────────────┘

Usage:
    python generate_pfq_doc.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "PFQ_Code_Changes.docx")

# ── Palette ────────────────────────────────────────────────────────────────────
RED         = RGBColor(0xC0, 0x00, 0x00)
GREEN       = RGBColor(0x00, 0x70, 0x00)
DARK_TEXT   = RGBColor(0x1F, 0x23, 0x28)
MUTED_TEXT  = RGBColor(0x6B, 0x72, 0x80)
CODE_BLUE   = RGBColor(0x17, 0x5B, 0xC2)
ACCENT_BLUE = RGBColor(0x2F, 0x81, 0xF7)

# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Solid background fill on a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:shd")):
        tcPr.remove(ex)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_border(cell, color="BBBBBB", sz="6"):
    """Box border on a single table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(ex)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        borders.append(b)
    tcPr.append(borders)


def _set_cell_no_border(cell):
    """Remove all borders from a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(ex)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        borders.append(b)
    tcPr.append(borders)


def _set_table_full_width(table):
    """Stretch a table to 100 % of the text area."""
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for ex in tblPr.findall(qn("w:tblW")):
        tblPr.remove(ex)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "5000")   # 5000 = 100 % in fiftieths-of-a-percent
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


def _set_cell_width_twips(cell, twips: int):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:tcW")):
        tcPr.remove(ex)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for ex in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(ex)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        borders.append(b)
    tblPr.append(borders)


def _para_spacing(para, before_pt=0, after_pt=0):
    pPr = para._p.get_or_add_pPr()
    for ex in pPr.findall(qn("w:spacing")):
        pPr.remove(ex)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(int(before_pt * 20)))
    sp.set(qn("w:after"),  str(int(after_pt  * 20)))
    pPr.append(sp)


# ── Document-level helpers ─────────────────────────────────────────────────────

def _add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before_pt=0, after_pt=2)
    r = p.add_run("PFQ — Code Changes Analysis")
    r.bold = True
    r.font.size  = Pt(18)
    r.font.color.rgb = ACCENT_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p2, before_pt=0, after_pt=14)
    r2 = p2.add_run("Security & Validation Review  |  IT Audit")
    r2.font.size  = Pt(10)
    r2.font.color.rgb = MUTED_TEXT


def _add_section_heading(doc, title: str):
    p = doc.add_paragraph()
    _para_spacing(p, before_pt=12, after_pt=4)
    r = p.add_run(title)
    r.bold = True
    r.font.size  = Pt(11)
    r.font.color.rgb = DARK_TEXT


# ── Code block (nested table) ──────────────────────────────────────────────────

def _add_code_block(cell, code_lines: list[str]):
    """
    Inserts a nested 1×1 table inside *cell* styled as a code block:
    Consolas 9pt, light-grey background, thin box border, no extra spacing.
    """
    nested = cell.add_table(rows=1, cols=1)
    _set_table_full_width(nested)
    _remove_table_borders(nested)

    cc = nested.cell(0, 0)
    _set_cell_bg(cc, "F2F4F6")
    _set_cell_border(cc, color="C0C0C0", sz="4")

    # Add a small top-padding paragraph (blank, 2 pt)
    first = True
    for line in code_lines:
        if first:
            cp = cc.paragraphs[0]
            cp.clear()
            first = False
        else:
            cp = cc.add_paragraph()
        run = cp.add_run(line)
        run.font.name  = "Consolas"
        run.font.size  = Pt(9)
        run.font.color.rgb = CODE_BLUE
        _para_spacing(cp, before_pt=0, after_pt=0)

    # Small bottom padding
    pad = cc.add_paragraph()
    _para_spacing(pad, before_pt=2, after_pt=0)


# ── Change-entry table ─────────────────────────────────────────────────────────
#
#  Column layout  (A4, margins 2.54 cm each side → 160 mm usable)
#    col 0  label   : 20 mm  →  1134 twips
#    col 1  content : 140 mm → 7938 twips
#
LABEL_TWIPS   = 1134   # ≈ 20 mm
CONTENT_TWIPS = 7938   # ≈ 140 mm


def _build_label_cell(cell, text: str, color: RGBColor, bg_hex: str):
    _set_cell_bg(cell, bg_hex)
    _set_cell_border(cell, "CCCCCC", "4")
    _set_cell_width_twips(cell, LABEL_TWIPS)
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before_pt=5, after_pt=5)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(9)
    r.font.color.rgb = color


def _build_content_cell(cell, description: str, code_lines: list[str] | None, bg_hex: str):
    _set_cell_bg(cell, bg_hex)
    _set_cell_border(cell, "CCCCCC", "4")
    _set_cell_width_twips(cell, CONTENT_TWIPS)

    p = cell.paragraphs[0]
    p.clear()
    _para_spacing(p, before_pt=4, after_pt=2 if not code_lines else 4)
    r = p.add_run(description)
    r.font.size  = Pt(10)
    r.font.color.rgb = DARK_TEXT

    if code_lines:
        _add_code_block(cell, code_lines)
        # small trailing spacer
        sp = cell.add_paragraph()
        _para_spacing(sp, before_pt=2, after_pt=0)


def add_change_entry(doc, change: dict):
    """
    Renders one full change entry table (file-path header + BEFORE + AFTER rows).
    change = {
        "file":   "/Path/To/File.cs",
        "before": { "description": "...", "code": None | ["line1", ...] },
        "after":  { "description": "...", "code": None | ["line1", ...] },
    }
    """
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    _set_table_full_width(table)

    # ── Row 0: file path (merged, full width) ──────────────────────────────────
    r0 = table.rows[0]
    r0.cells[0].merge(r0.cells[1])
    fc = r0.cells[0]
    _set_cell_bg(fc, "E8E8E8")
    _set_cell_border(fc, "AAAAAA", "8")
    p = fc.paragraphs[0]
    p.clear()
    _para_spacing(p, before_pt=5, after_pt=5)
    icon = p.add_run("📄  ")
    icon.font.size = Pt(9)
    run = p.add_run(change["file"])
    run.font.name  = "Consolas"
    run.font.size  = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)

    # ── Row 1: BEFORE ──────────────────────────────────────────────────────────
    _build_label_cell(
        table.rows[1].cells[0],
        "BEFORE", RED, "FFF0EE"
    )
    _build_content_cell(
        table.rows[1].cells[1],
        change["before"]["description"],
        change["before"].get("code"),
        "FAFAFA"
    )

    # ── Row 2: AFTER ───────────────────────────────────────────────────────────
    _build_label_cell(
        table.rows[2].cells[0],
        "AFTER", GREEN, "EEFBEE"
    )
    _build_content_cell(
        table.rows[2].cells[1],
        change["after"]["description"],
        change["after"].get("code"),
        "FAFAFA"
    )

    # Spacer after each entry
    sp = doc.add_paragraph()
    _para_spacing(sp, before_pt=0, after_pt=6)


# ══════════════════════════════════════════════════════════════════════════════
#  CONTENT — add as many entries as needed
# ══════════════════════════════════════════════════════════════════════════════

CHANGES = [
    {
        "file": "/WERCSmart.Portal.API.Supplier/Controllers/SupplierController.cs",
        "before": {
            "description": "Minimal validation for userB2CId.",
            "code": None,
        },
        "after": {
            "description": "Enhanced validation for userB2CId:",
            "code": [
                'if (userB2CId == Guid.Empty)',
                '{',
                '    _logger.LogWarning($"API-msg: New GetSupplier request. '
                'Invalid User B2C Id = {userB2CId}. ");',
                '    return BadRequest();',
                '}',
            ],
        },
    },
]


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    doc = Document()

    # Page margins (A4)
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    _add_title_block(doc)
    _add_section_heading(doc, "1.  Validation Changes")

    for change in CHANGES:
        add_change_entry(doc, change)

    doc.save(OUTPUT_PATH)
    size_kb = os.path.getsize(OUTPUT_PATH) // 1024
    print(f"Saved: {OUTPUT_PATH}  ({size_kb} KB)")
